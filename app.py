# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# LIBRARIES SECTION

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# Universally Unique Identifier lib for the creation of session and CAPTCHA
import uuid

# for managing the image paths
import os

# regex lib for regular expression check of input fields
import re

# provides necessary hash for the security of input data
import bcrypt

# Framework used : FLASK
from flask import Flask, redirect, request, render_template, session

# Flask 			: Flask class used to create web app
# redirect 			: redirecting to specific url
# request 			: get the input from HTML fields
# render_template 	: load specified HTML page
# session 			: maintain session and session variables

# create unique Sessions
from flask_sessionstore import Session

# create unique CAPTCHA for each session
from flask_session_captcha import FlaskSessionCaptcha

# schedule the update of attempts everyday
from flask_apscheduler import APScheduler

from flask_wtf.csrf import CSRFProtect

from flask_wtf.csrf import CSRFError

# connect MongoDB with Python
from pymongo import MongoClient

# added security for input files
from werkzeug.utils import secure_filename

# Keras for using the H5 model
from keras.models import Sequential
from keras.models import load_model
from keras.utils import np_utils

# Numpy to reshape and resize the pixel data of image
import numpy as np

# Computer Vision Library for object detection in image
import cv2

# Library for Pattern Matching
import glob

# PIL : Image porocessing library
from PIL import Image

# datetime to check whether the updates occur on time everday
# used for logging
from datetime import datetime as dt

# Library to check the file format by header
import magic

# Pandas used to create a dataframe of all pothole data
# the dataframe will be used to plot graph using plotly
import pandas as pd

# library used to plot interactive graphs for web-apps
import plotly

# used to plot the graph from pandas dataframe
import plotly.express as px

# .io used for input/output and save graph as image
import plotly.io as pio

# graph is converted to JSON which will be converted to interactive image
# via Javascript on admin dashboard
import json

# mail libraries
import email, smtplib, ssl

# used to format the sender account name
from email.utils import formataddr

# Encode file in ASCII characters to send by email
from email import encoders

# Multipurpose Internet Mail Extensions
# lets users exchange different kinds of data files, including audio, video, images and application programs, over email
from email.mime.base import MIMEBase

# used to create a multipart message
# text-content + mime attachment
from email.mime.multipart import MIMEMultipart

# add body text along with attachments
from email.mime.text import MIMEText

# library to work with Word Documents
import docx

# scale images in Word Documents
from docx.shared import Inches, Cm

# convert .docx to .pdf (report making)
from docx2pdf import convert

import string
import secrets
import smtplib
import ssl
from email.utils import formataddr
from email.message import EmailMessage
import os
import argon2 as ar


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# APP CONFIGURATIONS SECTION

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# path to uploads folder where necessary files will be stored
UPLOAD_FOLDER = "uploads\\"

# no other extensions allowed for uploaded file
ALLOWED_EXTENSIONS = {"jpg", "jpeg"}

# Flask Application Initialized
app = Flask(__name__)


# The Uploads Folder where the valid Images will be stored
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Uploaded file cannot be larger than 2 MB
app.config["MAX_CONTENT_LENGTH"] = 2 * 1024 * 1024

# MongoDB Client for connection
mongoClient = MongoClient("localhost", 27017)

# Key required for CAPTCHA
app.config["SECRET_KEY"] = uuid.uuid4()

app.config["WTF_CSRF_SECRET_KEY"] = str(uuid.uuid4())

app.config["CAPTCHA_ENABLE"] = True

# Count of Numbers in the CAPTCHA
app.config["CAPTCHA_LENGTH"] = 6

# CAPTCHA Image Dimensions
app.config["CAPTCHA_WIDTH"] = 160
app.config["CAPTCHA_HEIGHT"] = 60

# MongoDB Session Parameters used for generating CAPTCHA
app.config["SESSION_MONGODB"] = mongoClient
app.config["SESSION_TYPE"] = "mongodb"
app.config["SESSION_PERMANENT"] = False


# Flask Session for creating CAPTCHA
Session(app)

# CAPTCHA created using Flask Session
captcha = FlaskSessionCaptcha(app)

csrf = CSRFProtect(app)

# scheduler objects to schedule tasks

# reset upload count everyday at midnight
file_upload_attempt_update = APScheduler()

# reporting authorities
authority_report = APScheduler()

# clear previous sessions every 6 hours
clr_sessions = APScheduler()

# remove unverified accounts
clr_unverified_accounts = APScheduler()

# reset verification code count everyday at midnight
verification_code_attempt_update = APScheduler()

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# SCHEDULED FUNCTIONS

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# Function to update attempts everyday at
def update_attempts():
    # the database that stores user data
    db = mongoClient["userdata"]

    # collection that stores user data
    collection = db["data"]

    # the update query that updates the attempts
    # count to 3 everyday
    if collection.update_many({}, {"$set": {"attempts": 3}}):
        # count of all the updated documents
        n = collection.count_documents({"attempts": 3})

        # total documents in the collection
        t = collection.count_documents({})

        # printing the counts to ensure all records are updated
        print(f"==> {dt.now()} :: {n}/{t} Documents Updated")


#
def clear_sessions():
    # the database that stores sessions
    db = mongoClient["flask_sessionstore"]

    # collection that stores sessions
    collection = db["sessions"]

    # delete all previous sessions
    collection.drop()
    print(f"==> {dt.now()} :: All sessions cleared")


def update_verification_code_attempts():
    # the database that stores user data
    db = mongoClient["userdata"]

    # collection that stores user data
    collection = db["code_count"]

    collection.drop()

    collection = db["data"]

    all_doc = collection.find({}, {"_id": 0, "mail": 1})

    ls = [value for doc in all_doc for value in doc.values()]

    print(ls)

    collection = db["code_count"]

    for i in ls:
        collection.insert_one({"mail": i, "code": "0", "attempts": 3})

    # count of all the updated documents
    n = collection.count_documents({"code": "0", "attempts": 3})

    # total documents in the collection
    t = collection.count_documents({})

    # printing the counts to ensure all records are updated
    print(f"==> {dt.now()} :: {n}/{t} Documents Updated")


def clear_unverified_accounts():
    # the database that stores user data
    db = mongoClient["userdata"]

    # collection that stores user data
    collection = db["data"]

    # the update query that updates the attempts
    # count to 3 everyday
    if collection.delete_many({"verified": 0}):
        # count of all the updated documents
        n = collection.count_documents({"verified": 1})

        # total documents in the collection
        t = collection.count_documents({})

        # printing the counts to ensure all records are updated
        print(f"==> {dt.now()} :: {n}/{t} Documents Remain")


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# FILE-RELATED FUNCTIONS

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# Function to check the validity of filename
def allowed_file(filename):
    file_regex = r"^[a-zA-Z0-9\-\_]{1,32}\.(jpeg|jpg)$"
    return (
        "." in filename
        and filename.rsplit(".", 1)[-1].lower() in ALLOWED_EXTENSIONS
        and re.fullmatch(file_regex, filename)
    )


# check file validity from the file header
def valid_file(file):
    #
    ls = str(magic.from_file(file, mime=True)).lower().split("/")
    print("-> ", ls)
    print(ls[0], ls[1])
    return ls[1] in ALLOWED_EXTENSIONS and ls[0] == "image"
    # return 'image' in magic.from_file(file, mime=True)


# Machine Learning Model for Pothole Detection
def model(image):
    global result, pothole_percent, size

    size = 300
    model = Sequential()
    model = load_model("full_model.h5")

    test = [cv2.imread(img, 0) for img in image]
    print(test)

    for i in range(len(test)):
        test[i] = cv2.resize(test[i], (size, size))

    X_temp = np.asarray(test)

    X_test = []
    X_test.extend(X_temp)
    X_test = np.asarray(X_test)
    X_test = X_test.reshape(X_test.shape[0], size, size, 1)

    Y_temp = np.zeros([X_temp.shape[0]], dtype=int)
    # Y_temp = np.ones([X_temp.shape[0]], dtype = int)

    Y_test = []
    Y_test.extend(Y_temp)
    Y_test = np.asarray(Y_test)
    Y_test = np_utils.to_categorical(Y_test)

    X_test = X_test / 255

    tests = model.predict(X_test)

    for i in range(len(X_test)):
        print(">>> Predicted %d = %s" % (i, tests[i]))

    result = tests[i]
    pothole_percent = float("{0:.2f}".format(result[1] * 100))

    detect_result = "Pothole Detected"

    print(result)
    if result[1] < 0.6:
        detect_result = "No " + detect_result

    return pothole_percent, detect_result


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# DATABASE FUNCTIONS

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# Function to decrement attempt after uploading image
def decrement_attempt(mail):
    # the database that stores user data
    db = mongoClient["userdata"]

    # collection that stores user data
    collection = db["data"]

    # the mail will be used as query for decrementing
    # the upload attempts count
    details = {"mail": "%s" % (mail)}

    # check if the document with the required mail
    # attempts gets decremented by 1
    # every time a file is successfully uploaded
    if collection.update_many(details, {"$inc": {"attempts": -1}}):
        # fetching the document with the mail
        doc = list(collection.find({"mail": "%s" % (mail)}, {}))

        # print the document to verify
        # successful decrement of attempt count
        print(f"==> {doc}")


# Function to return remaining attempts
def get_attempts(mail):
    # the database that stores user data
    db = mongoClient["userdata"]

    # collection that stores user data
    collection = db["data"]

    # the mail will be used as query for fetching
    # the upload attempts count
    details = {"mail": "%s" % (mail)}

    # list with mail as query
    # returns only the attempt count in BSON
    attempts = list(collection.find(details, {"attempts": 1}))

    # getting the attempt count from the BSON
    attempt = attempts[0]["attempts"]
    print(attempt, type(attempt))

    # the attempt count is returned
    return attempt


# Function to return remaining attempts
def get_code_attempts(mail):
    # the database that stores user data
    db = mongoClient["userdata"]

    # collection that stores user data
    collection = db["code_count"]

    # the mail will be used as query for fetching
    # the upload attempts count
    details = {"mail": "%s" % (mail)}

    # list with mail as query
    # returns only the attempt count in BSON
    attempts = list(collection.find(details, {"attempts": 1}))

    # getting the attempt count from the BSON
    attempt = attempts[0]["attempts"]
    print(attempt, type(attempt))

    # the attempt count is returned
    return attempt


# Function to return remaining attempts
def get_code(mail):
    # the database that stores user data
    db = mongoClient["userdata"]

    # collection that stores user data
    collection = db["code_count"]

    # the mail will be used as query for fetching
    # the upload attempts count
    details = {"mail": "%s" % (mail)}

    # list with mail as query
    # returns only the attempt count in BSON
    codes = list(collection.find(details, {"code": 1}))

    # getting the attempt count from the BSON
    code = codes[0]["code"]
    print(code, type(code))

    # the attempt count is returned
    return code


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# REPORT MAIL FUNCTION

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# Function to create PDF Report and Mail to respective authorities
def report_authority():
    bar_state, bar_state_district = get_counts(True)

    doc = docx.Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    subject = "Weekly Report-" + dt.now().strftime("%d-%m-%Y")

    # Add a Title to the document
    doc.add_heading(subject, 0)

    doc.add_heading("State-wise Pothole Reports", 1)

    # Image with defined size
    doc.add_picture("state_bar_graph.jpg", width=Cm(16), height=Cm(12))

    doc.add_page_break()

    doc.add_heading("District-wise Pothole Reports", 1)

    doc.add_picture("state_district_bar_graph.jpg", width=Cm(16), height=Cm(12))

    doc.add_page_break()
    # doc.add_picture('graph.png', width=Cm(15), height=Cm(11))

    doc.add_heading("All Records", 0)

    table = doc.add_table(rows=1, cols=8, style="Table Grid")

    row = table.rows[0].cells

    head = [
        "address",
        "area",
        "state",
        "district",
        "zipcode",
        "file_name",
        "pothole_percent",
        "timestamp",
    ]

    for i, j in enumerate(head):
        row[i].text = j.capitalize()

    # client = mongoClient(host="127.0.0.1", port=27017)

    db = mongoClient["userdata"]

    collection = db["pothole_result"]

    all_doc = list(collection.find({}, {"_id": 0, "mail": 0}))

    ls = [[value for value in doc.values()] for doc in all_doc]

    df = pd.DataFrame.from_records(list(all_doc))

    print(df)

    csv_file = "Weekly Report-" + dt.now().strftime("%d-%m-%Y") + ".csv"

    df.to_csv(csv_file, index=False)

    for data in ls:
        row = table.add_row().cells
        for i, j in enumerate(data):
            row[i].text = j

    table.style = "Colorful List"
    # table.style = "Light List"

    # Now save the document to a location
    file = subject + ".docx"

    # doc.add_page_break()

    # images = os.listdir(UPLOAD_FOLDER)

    # print(images)
    # for img in images:
    #     doc.add_paragraph(img)
    #     print(UPLOAD_FOLDER + img)
    #     doc.add_picture("/uploads/" + img, width=Cm(8), height=Cm(6))

    doc.save(file)
    print("docx saved")
    print([i for i in os.listdir() if ".docx" in i])

    # convert(file)
    # print("pdf saved")
    # print([i for i in os.listdir() if ".pdf" in i])

    # os.remove(file)
    os.remove("state_bar_graph.jpg")
    os.remove("state_district_bar_graph.jpg")

    db = mongoClient["userdata"]

    collection = db["admin_mail"]

    all_doc = collection.find({}, {"_id": 0})

    receiver_emails = [value for doc in all_doc for value in doc.values()]

    # ls= [i for i in j for j in ls]

    # print(receiver_emails)

    body = (
        "Road Health Management Report "
        + dt.now().strftime("%d-%m-%Y")
        + " attached below PFA."
    )
    sender_email = "anand.extra.01@gmail.com"
    receiver_email = "omanand2002@gmail.com"
    password = os.environ.get("extra01_mail_python")

    subject = "Weekly Report-" + dt.now().strftime("%d-%m-%Y")

    # Create a multipart message and set headers
    message = MIMEMultipart()
    message["From"] = formataddr(("Road Health Management", sender_email))
    message["To"] = receiver_email
    # message["To"] = ",".join(m for m in receiver_emails)
    message["Subject"] = subject
    # message["Bcc"] = receiver_email  # Recommended for mass emails

    # Add body to email
    message.attach(MIMEText(body, "plain"))

    files = [
        "Weekly Report-" + dt.now().strftime("%d-%m-%Y") + ".docx",
        "Weekly Report-" + dt.now().strftime("%d-%m-%Y") + ".csv",
    ]  # In same directory as script

    for filename in files:
        # Open PDF file in binary mode
        with open(filename, "rb") as attachment:
            # Add file as application/octet-stream
            # Email client can usually download this automatically as attachment
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())

        # Encode file in ASCII characters to send by email
        encoders.encode_base64(part)

        # Add header as key/value pair to attachment part
        part.add_header(
            "Content-Disposition",
            f"attachment; filename= {filename}",
        )

        # Add attachment to message and convert message to string
        message.attach(part)

    text = message.as_string()

    # Log in to server using secure context and send email
    context = ssl.create_default_context()
    with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
        server.login(sender_email, password)
        # server.sendmail(sender_email, receiver_emails, text)
        server.sendmail(sender_email, receiver_email, text)

    print(f"==> {dt.now()} :: Emails sent")


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# GRAPH FUNCTIONS

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


def graph_state(count_by_state, save_graph):
    state_df = pd.DataFrame(count_by_state, columns=["State", "Potholes"])

    print(state_df)

    # Create Bar chart
    fig = px.bar(
        state_df,
        y="Potholes",
        x="State",
        color="State",
        title="State-wise Pothole Records",
    )

    # Create graphJSON
    graph_bar = json.dumps(fig, cls=plotly.utils.PlotlyJSONEncoder)

    if save_graph:
        pio.write_image(fig, "state_bar_graph.jpg")

    return graph_bar


def graph_state_district(count_by_state_district, save_graph):
    district_state_df = pd.DataFrame(
        count_by_state_district, columns=["District", "Potholes", "State"]
    )
    # print(district_state_df)

    # Create Bar chart
    fig = px.bar(
        district_state_df,
        y="Potholes",
        x="District",
        color="State",
        title="District-wise Pothole Records",
    )

    # Create graphJSON
    graph_bar = json.dumps(fig, cls=plotly.utils.PlotlyJSONEncoder)

    if save_graph:
        pio.write_image(fig, "state_district_bar_graph.jpg")

    return graph_bar


def get_counts(save_graph=False):
    db = mongoClient["userdata"]

    collection = db["pothole_result"]

    # STATES_LIST_INDIA = ["Gujarat", "Maharashtra", "Delhi"]

    # print(collection.distinct("state"))
    # print(collection.distinct("district", {"state": "Gujarat"}))

    states = collection.distinct("state")

    count_by_state = []

    count_by_state_district = []

    for state in states:
        state_count = collection.count_documents({"state": state})
        # print(state, state_count)
        count_by_state.append([state, state_count])
        districts = collection.distinct("district", {"state": state})
        for district in districts:
            district_count = collection.count_documents({"district": district})
            # print(state, district, district_count)
            count_by_state_district.append([district, district_count, state])

    # print(count_by_state)
    # print(count_by_state_district)

    return (
        graph_state(count_by_state, save_graph),
        graph_state_district(count_by_state_district, save_graph),
    )


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# FLASK ROUTES AND FUNCTIONS

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# Default Index Route
@app.route("/")
def index():
    session["mail"] = None
    session["password_recovery"] = False
    return render_template("index.html")


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# Login Page Function
@app.route("/login")
def login():
    check = {
        "valid_Mail": True,
        "valid_Pass": True,
        "valid_user": True,
        "valid_Captcha": True,
    }
    msg = ""
    values = {"mail": ""}
    session["mail"] = None
    # the checks are by default all True
    # After validation and db checking, the checks can be False
    # empty msg as the form is not filled
    # empty value for mail as first time the form is not filled
    # (first time access of form)

    # login.html rendered with default values
    return render_template("login.html", check=check, msg=msg, values=values)


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# Signup Page function
@app.route("/signup")
def signup():
    check = {
        "valid_Mail": True,
        "mail_Duplicate": True,
        "valid_Pass": True,
        "match_Pass": True,
        "valid_Captcha": True,
    }
    msg = ""
    values = {"mail": ""}
    session["mail"] = None
    # the checks are by default all True
    # After validation and db checking, the checks can be False
    # empty msg as the form is not filled
    # empty value for mail as first time the form is not filled
    # (first time access of form)

    # signup.html rendered with default values
    return render_template("signup.html", check=check, msg=msg, values=values)


@app.route("/forgot_password")
def forgot_password():
    check = {"valid_Mail": True, "user_exist": True}
    values = {"mail": ""}
    return render_template("password_recovery.html", check=check, values=values)


@app.route("/check_mail", methods=["POST", "GET"])
def check_mail():
    check = {"valid_Mail": True, "user_exist": True}
    mail_regex = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b"

    db = mongoClient["userdata"]

    # the collection used for valid user credentials
    collection = db["data"]

    if request.method == "POST":
        tmp_mail = request.form.get("mail").strip().replace(" ", "")

        session["tmp_mail"] = tmp_mail

        values = {"mail": tmp_mail}

        if not tmp_mail or not re.fullmatch(mail_regex, tmp_mail):
            check["valid_Mail"] = False
            return render_template("password_recovery.html", check=check, values=values)

        else:
            if collection.count_documents({"mail": "%s" % (tmp_mail)}):
                # if user already exists

                if collection.count_documents(
                    {"mail": "%s" % (tmp_mail), "verified": 1}
                ):
                    # user exists and is verified
                    session["password_recovery"] = True
                    return redirect("/code_generate")

                else:
                    # user unverified
                    return redirect("/code_generate")

                # collection = db["code_count"]

                # if collection.count_documents

            else:
                # no such user exist
                check["user_exist"] = False
                return render_template(
                    "password_recovery.html", check=check, values=values
                )
    return render_template("password_recovery.html", check=check, values=values)


@app.route("/code_generate", methods=["POST", "GET"])
def code_gen():
    code = "".join(secrets.choice(string.digits) for _ in range(6))
    check = {"valid_code": True, "code_attempt": True}

    db = mongoClient["userdata"]

    # collection that stores user data
    collection = db["code_count"]

    if get_code_attempts(session["tmp_mail"]) > 0:
        collection.update_one(
            {"mail": session["tmp_mail"]},
            {"$set": {"code": code}, "$inc": {"attempts": -1}},
        )

        email_sender = "anand.extra.01@gmail.com"
        email_password = os.environ.get("extra01_mail_python")
        email_receiver = session["tmp_mail"]

        subject = "Verification Code"
        body = "Verification Code : " + code

        em = EmailMessage()
        em["From"] = formataddr(("Road Health Management", email_sender))
        em["To"] = email_receiver
        em["Subject"] = subject

        em.set_content(body)
        context = ssl.create_default_context()

        with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as smtp:
            smtp.login(email_sender, email_password)
            smtp.sendmail(email_sender, email_receiver, em.as_string())

        print(code)
    else:
        check["code_attempt"] = False

    return render_template(
        "code_verify.html", check=check, user_verified=session["password_recovery"]
    )


@app.route("/code_verify", methods=["POST", "GET"])
def code_verify():
    code = get_code(session["tmp_mail"])
    print(code)

    check = {"valid_code": True, "code_attempt": True}

    user_code = request.form.get("code").strip().replace(" ", "")

    print(session["tmp_mail"], session["password_recovery"])
    if user_code == code:
        if session["password_recovery"]:
            return redirect("/change_password")
        else:
            db = mongoClient["userdata"]
            collection = db["data"]

            collection.update_one(
                {"mail": session["tmp_mail"]},
                {"$set": {"verified": 1}},
            )

            print(list(db["data"].find({"mail": session["tmp_mail"]})))

            return render_template(
                "verified_user.html", user_verified=True, pass_change=False
            )
    else:
        check["valid_code"] = False

    return render_template(
        "code_verify.html", check=check, user_verified=session["password_recovery"]
    )


@app.route("/change_password", methods=["POST", "GET"])
def change_password():
    check = {
        "valid_Pass": True,
        "match_Pass": True,
    }
    return render_template("change_password.html", check=check)


@app.route("/check_password", methods=["POST", "GET"])
def check_password():
    check = {
        "valid_Pass": True,
        "match_Pass": True,
    }

    password = request.form.get("password").strip().replace(" ", "")
    conf_pass = request.form.get("conf_pass").strip().replace(" ", "")

    pass_regex = (
        r"^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*#?&])[A-Za-z\d@$!#%*?&]{8,20}$"
    )

    if (
        not password
        or not conf_pass
        or not re.fullmatch(pass_regex, password)
        or not re.fullmatch(pass_regex, conf_pass)
    ):
        check["valid_Pass"] = False

    if password != conf_pass:
        check["match_Pass"] = False

    if all([i for i in check.values()]):
        mail = session["tmp_mail"]
        # salt = bcrypt.gensalt()
        # password = bcrypt.hashpw(password.encode(), salt)

        ph = ar.PasswordHasher(hash_len=64)

        pepper1, pepper2 = mail.split("@")

        password = ph.hash(pepper1 + password + pepper2)

        print(pepper1)
        print(pepper2)
        print(password)

        db = mongoClient["userdata"]
        collection = db["data"]

        collection.update_one(
            {"mail": mail},
            {"$set": {"pass": "%s" % password}},
        )

        print("password changed successfully", password)

        print(list(db["data"].find({"mail": session["tmp_mail"]})))

        return render_template(
            "verified_user.html", user_verified=False, pass_change=True
        )

    return render_template("change_password.html", check=check)


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# I M P O R T A N T
# 					I M P O R T A N T
# 									  I M P O R T A N T


# function to validate and check the login and signup information
@app.route("/form", methods=["POST", "GET"])
def form():
    file_check = {
        "valid_address": True,
        "valid_area": True,
        "valid_state": True,
        "valid_district": True,
        "valid_zip": True,
        "valid_file": True,
        "valid_Captcha": True,
    }
    file_values = {"address": "", "area": "", "zipcode": ""}
    # all the file upload form checks are kept True by default
    # all file_values are kept empty before the form is filled for the first time
    # (first time access of form)

    # database to be used for verifying Login info
    # or add the credentials after Signup
    db = mongoClient["userdata"]

    # the collection used for valid user credentials
    collection = db["data"]

    mail_regex = r"^[a-zA-Z0-9.]{1,64}@(?:[a-zA-Z0-9-]{1,63}\.){1,8}[a-zA-Z]{2,63}$"
    # Regular Expression for valid mail id pattern
    # --> does not check the validity(existence) of mail <--

    pass_regex = (
        r"^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*#?&])[A-Za-z\d@$!#%*?&]{8,20}$"
    )
    # Regular Expression for valid password
    # must include :
    # 1 Uppercase
    # 1 Lowercase
    # 1 Number/Digit
    # 1 Special Character
    # length must be between 8 to 20

    # check for session started
    # mail = session.get('mail')
    # print(mail)

    # Login/Signup form redirect
    # All POST requests get validated here ↓
    if request.method == "POST":
        # Condition for Signup Form
        if request.form.get("page") == "Signup":
            check = {
                "valid_Mail": True,
                "mail_Duplicate": True,
                "valid_Pass": True,
                "match_Pass": True,
                "valid_Captcha": True,
            }
            # checks will be determined after checking all fields against regular expressions and DB

            # field values given by the user for account creation
            mail = request.form.get("mail").strip().replace(" ", "")
            password = request.form.get("password").strip().replace(" ", "")
            conf_pass = (
                request.form.get("conf_pass").strip().replace(" ", "")
            )  # confirm password field
            # the password & confirm password are received in plain-text
            # the leading, trailing & in-between spaces are removed by strip() & replace()

            # will be used as value in case of unsuccessful signup
            values = {"mail": mail}

            # check for mail
            # should not be empty
            # match the regular Expression
            if not mail or not re.fullmatch(mail_regex, mail):
                # if any of the condition fails the
                # valid mail check is marked False
                check["valid_Mail"] = False

            # checking if the mail already exists in the DB
            if collection.count_documents({"mail": "%s" % (mail)}):
                # if user already exists
                check["mail_Duplicate"] = False

            # check for plain-text password & confirm password
            # should not be empty
            # match the regular Expression
            if (
                not password
                or not conf_pass
                or not re.fullmatch(pass_regex, password)
                or not re.fullmatch(pass_regex, conf_pass)
            ):
                # if any of the condition fails the
                # valid password check is marked False
                check["valid_Pass"] = False

            # match password and confirm password fields
            if password != conf_pass:
                # if the passwords do not match
                # match password check is marked False
                check["match_Pass"] = False

            # validation of CAPTCHA
            if not captcha.validate():
                # if incorrect captcha
                check["valid_Captcha"] = False

            # generating random salt that will be used to
            # add extra security after hashing of plaintext password
            # salt type -> bytes
            salt = bcrypt.gensalt()

            # plain-text password is encoded to bytes and
            # salted to get hashed password
            # hashed password type -> bytes
            password = bcrypt.hashpw(password.encode(), salt)

            # if all the checks are True
            if all([i for i in check.values()]):
                # dictionary used as query for creating user in DB
                # and return same values to form in case there is an error
                details = {
                    "mail": "%s" % (mail),
                    "pass": "%s" % (password.decode()),
                    "attempts": 3,
                    "verified": 0,
                }

                # inersting the details in DB
                # creating the user account
                if collection.insert_one(details):
                    # mail id will be stored in token session
                    # to keep track of file uploads
                    session["tmp_mail"] = mail

                    collection = db["code_count"]

                    collection.insert_one({"mail": mail, "code": "0", "attempts": 3})

                    # initially declaration of remaining attempts
                    remaining_attempts = False

                    # the actual count of attempts remaining based on mail
                    attempts = get_attempts(mail)

                    # mark remaining attempts as True if user has got attempts remaining
                    if attempts > 0:
                        remaining_attempts = True

                    return redirect("/code_generate")
                    # redirected to file upload form after successful user login
                    # return render_template(
                    #     "file_upload.html",
                    #     msg="User Created Successfully",
                    #     check=file_check,
                    #     err="",
                    #     attempts=attempts,
                    #     remaining_attempts=remaining_attempts,
                    #     values=file_values,
                    # )

                # incase there is an error in user creation
                # regarding DB
                else:
                    # redirect to signup page if there was error in user creation
                    return render_template(
                        "signup.html",
                        check=check,
                        msg="Failed to create user. Try Again",
                        values=values,
                    )

            # in case there is an error in input data
            # like mail/password regex error
            else:
                # # check if the session is active
                if session["mail"]:
                    # initially declaration of remaining attempts
                    remaining_attempts = False

                    # the actual count of attempts remaining based on mail
                    attempts = get_attempts(mail)

                    # mark remaining attempts as True if user has got attempts remaining
                    if attempts > 0:
                        remaining_attempts = True

                    # redirected to file upload form if form is refershed without submitting
                    return render_template(
                        "file_upload.html",
                        check=file_check,
                        msg="",
                        err="",
                        attempts=attempts,
                        remaining_attempts=remaining_attempts,
                        values=file_values,
                    )

                # incase session is not active
                else:
                    # the credentials are either empty or do not match the regular expression check
                    return render_template(
                        "signup.html", check=check, msg="Try Again", values=values
                    )

        # - - - - - - - - - - - - - - - - - - - - - - - - - - -

        # Condition for Login Form
        elif request.form.get("page") == "Login":
            check = {
                "valid_Mail": True,
                "valid_Pass": True,
                "valid_user": True,
                "valid_Captcha": True,
            }
            # checks will be determined after checking all fields against regular expressions and DB

            # the input credentials for Login
            mail = request.form.get("mail").strip().replace(" ", "")
            password = request.form.get("password").strip().replace(" ", "")
            # the password is received in plain-text
            # the leading, trailing & in-between spaces are removed by strip() & replace()

            # will be used as value in case of unsuccessful signup
            values = {"mail": mail}

            # check for mail
            # should not be empty
            # match the regular Expression
            if not mail or not re.fullmatch(mail_regex, mail):
                # if any of the condition fails the
                # valid mail check is marked False
                check["valid_Mail"] = False

            # check for plain-text password
            # should not be empty
            # match the regular Expression
            if not password or not re.fullmatch(pass_regex, password):
                # if any of the condition fails the
                # valid password check is marked False
                check["valid_Pass"] = False

            # validation of CAPTCHA
            if not captcha.validate():
                # if incorrect captcha
                check["valid_Captcha"] = False

            print(check)
            print(session["mail"])
            # if all the checks are True
            if all([i for i in check.values()]):
                # the count of user will be returned (0 or 1)
                # 1 if exists ; 0 if not
                # the DB contains only unique mails
                user_exist = collection.count_documents({"mail": "%s" % (mail)})
                print(user_exist)

                # user found in DB
                if user_exist:
                    if collection.count_documents(
                        {"mail": "%s" % (mail), "verified": 1}
                    ):
                        # the password field will be returned for the mail id entered
                        db_passwords = list(
                            collection.find({"mail": "%s" % (mail)}, {"pass": 1})
                        )
                        db_password = db_passwords[0]["pass"]

                        # password hashed and salted using Bcrypt
                        # checkpw function used to check the input password
                        # the function takes password(input) & password stored in database
                        # both parameters need to be in bytes and thus are encoded
                        # boolean value is returned (True for match)
                        correct_pass = bcrypt.checkpw(
                            password.encode(), db_password.encode()
                        )

                        print(user_exist)
                        print(password)
                        print(db_password)
                        print(correct_pass)
                        print(user_exist and correct_pass)

                        # if the user is found and password is correct
                        if user_exist and correct_pass:
                            # mail id will be stored in token session
                            # to keep track of file uploads
                            session["mail"] = mail
                            session["tmp_mail"] = None

                            # initially declaration of remaining attempts
                            remaining_attempts = False

                            # the actual count of attempts remaining based on mail
                            attempts = get_attempts(mail)

                            # mark remaining attempts as True if user has got attempts remaining
                            if attempts > 0:
                                remaining_attempts = True

                            # redirected to file upload form after successful user login
                            return render_template(
                                "file_upload.html",
                                msg="Logged in Successfully",
                                check=file_check,
                                err="",
                                attempts=attempts,
                                remaining_attempts=remaining_attempts,
                                values=file_values,
                            )

                        else:
                            check["valid_Pass"] = False
                            check["valid_Mail"] = False
                            # redirect to login page if the user does not exist in DB
                            return render_template(
                                "login.html",
                                check=check,
                                msg="Incorrect Password",
                                values=values,
                            )
                    else:
                        return redirect("/code_generate")

                        # redirect to login page if the user does not exist in DB
                        # return render_template(
                        #     "login.html",
                        #     check=check,
                        #     msg="User does not exist.",
                        #     values=values,
                        # )
                else:
                    # redirect to login page if the user does not exist in DB
                    print("user not exist")
                    check["valid_user"] = False
                    return render_template(
                        "login.html",
                        check=check,
                        msg="User does not exist.",
                        values=values,
                    )

            # in case there is an error in input data
            # like mail/password regex error
            else:
                # # check if the session is active
                if session["mail"]:
                    # initially declaration of remaining attempts
                    remaining_attempts = False

                    # the actual count of attempts remaining based on mail
                    attempts = get_attempts(mail)

                    # mark remaining attempts as True if user has got attempts remaining
                    if attempts > 0:
                        remaining_attempts = True

                    # redirected to file upload form if form is refershed without submitting
                    return render_template(
                        "file_upload.html",
                        check=file_check,
                        msg="",
                        err="",
                        attempts=attempts,
                        remaining_attempts=remaining_attempts,
                        values=file_values,
                    )

                else:
                    # the credentials are either empty or do not match the regular expression check
                    check["valid_Pass"] = False
                    check["valid_user"] = False
                    check["valid_Mail"] = False
                    return render_template(
                        "login.html", check=check, msg="Try Again", values=values
                    )

    # - - - - - - - - - - - - - - - - - - - - - - - - - - -

    # elif mail:

    # 	print('elif', mail)
    # 	remaining_attempts= False

    # 	attempts= get_attempts(mail)

    # 	if attempts> 0:
    # 		remaining_attempts= True

    # 	return render_template("file_upload.html", check= check, msg= '', err= '', attempts= attempts, remaining_attempts= remaining_attempts)

    else:
        # print("redirct")
        # if the page is accesses without logging in or signing up
        # redirected to login page
        return redirect("/login")


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


@app.route("/login_validate", methods=["POST", "GET"])
def login_validate():
    file_check = {
        "valid_address": True,
        "valid_area": True,
        "valid_state": True,
        "valid_district": True,
        "valid_zip": True,
        "valid_file": True,
        "valid_Captcha": True,
    }
    msg = ""
    file_values = {"address": "", "area": "", "zipcode": ""}

    db = mongoClient["userdata"]

    collection = db["data"]

    mail_regex = r"^[a-zA-Z0-9.]{1,64}@(?:[a-zA-Z0-9-]{1,63}\.){1,8}[a-zA-Z]{2,63}$"

    pass_regex = (
        r"^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*#?&])[A-Za-z\d@$!#%*?&]{8,20}$"
    )

    check = {
        "valid_Mail": True,
        "valid_Pass": True,
        "valid_user": True,
        "valid_Captcha": True,
    }
    # checks will be determined after checking all fields against regular expressions and DB

    if request.method == "POST":
        # the input credentials for Login
        mail = request.form.get("mail").strip().replace(" ", "")
        password = request.form.get("password").strip().replace(" ", "")
        # the password is received in plain-text
        # the leading, trailing & in-between spaces are removed by strip() & replace()

        # will be used as value in case of unsuccessful signup
        values = {"mail": mail}

        # check for mail
        # should not be empty
        # match the regular Expression
        if not mail or not re.fullmatch(mail_regex, mail):
            # if any of the condition fails the
            # valid mail check is marked False
            check["valid_Mail"] = False

        # check for plain-text password
        # should not be empty
        # match the regular Expression
        if not password or not re.fullmatch(pass_regex, password):
            # if any of the condition fails the
            # valid password check is marked False
            check["valid_Pass"] = False

        # validation of CAPTCHA
        if not captcha.validate():
            # if incorrect captcha
            check["valid_Captcha"] = False

        print(check)
        print(session["mail"])

        # if all the checks are True
        if all([i for i in check.values()]):
            # the count of user will be returned (0 or 1)
            # 1 if exists ; 0 if not
            # the DB contains only unique mails
            user_exist = collection.count_documents({"mail": "%s" % (mail)})
            print(user_exist)

            # user found in DB
            if user_exist:
                if collection.count_documents({"mail": "%s" % (mail), "verified": 1}):
                    # the password field will be returned for the mail id entered
                    db_passwords = list(
                        collection.find({"mail": "%s" % (mail)}, {"pass": 1})
                    )
                    db_password = db_passwords[0]["pass"]

                    # password hashed and salted using Bcrypt
                    # checkpw function used to check the input password
                    # the function takes password(input) & password stored in database
                    # both parameters need to be in bytes and thus are encoded
                    # boolean value is returned (True for match)
                    # correct_pass = bcrypt.checkpw(
                    #     password.encode(), db_password.encode()
                    # )

                    ph = ar.PasswordHasher(hash_len=64)

                    pepper1, pepper2 = mail.split("@")

                    # password = ph.hash(pepper1 + password + pepper2)

                    password = pepper1 + password + pepper2

                    print(pepper1)
                    print(pepper2)
                    print(password)

                    correct_pass = True

                    try:
                        ph.verify(db_password, password)
                    except ar.exceptions.VerifyMismatchError:
                        correct_pass = False
                        print("The password does not match the supplied hash")

                    print(user_exist)
                    print(password)
                    print(db_password)
                    print(correct_pass)
                    print(user_exist and correct_pass)

                    # if the user is found and password is correct
                    if user_exist and correct_pass:
                        # mail id will be stored in token session
                        # to keep track of file uploads
                        session["mail"] = mail
                        session["tmp_mail"] = None

                        # initially declaration of remaining attempts
                        remaining_attempts = False

                        # the actual count of attempts remaining based on mail
                        attempts = get_attempts(mail)

                        # mark remaining attempts as True if user has got attempts remaining
                        if attempts > 0:
                            remaining_attempts = True

                        return redirect("/file_upload_form")

                        # redirected to file upload form after successful user login
                        return render_template(
                            "file_upload.html",
                            msg="Logged in Successfully",
                            check=file_check,
                            err="",
                            attempts=attempts,
                            remaining_attempts=remaining_attempts,
                            values=file_values,
                        )

                    else:
                        check["valid_Pass"] = False
                        check["valid_Mail"] = False
                        # redirect to login page if the user does not exist in DB
                        return render_template(
                            "login.html", check=check, msg=msg, values=values
                        )
                else:
                    return redirect("/code_generate")

                    # redirect to login page if the user does not exist in DB
                    # return render_template(
                    #     "login.html",
                    #     check=check,
                    #     msg="User does not exist.",
                    #     values=values,
                    # )
            else:
                # redirect to login page if the user does not exist in DB
                print("user not exist")
                check["valid_user"] = False
                return render_template(
                    "login.html", check=check, msg=msg, values=values
                )

        # in case there is an error in input data
        # like mail/password regex error
        else:
            return render_template("login.html", check=check, msg=msg, values=values)

    return redirect("/login")


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


@app.route("/signup_validate", methods=["POST", "GET"])
def signup_validate():
    file_check = {
        "valid_address": True,
        "valid_area": True,
        "valid_state": True,
        "valid_district": True,
        "valid_zip": True,
        "valid_file": True,
        "valid_Captcha": True,
    }
    msg = ""
    file_values = {"address": "", "area": "", "zipcode": ""}

    db = mongoClient["userdata"]

    collection = db["data"]

    mail_regex = r"^[a-zA-Z0-9.]{1,64}@(?:[a-zA-Z0-9-]{1,63}\.){1,8}[a-zA-Z]{2,63}$"

    pass_regex = (
        r"^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*#?&])[A-Za-z\d@$!#%*?&]{8,20}$"
    )

    if request.method == "POST":
        check = {
            "valid_Mail": True,
            "mail_Duplicate": True,
            "valid_Pass": True,
            "match_Pass": True,
            "valid_Captcha": True,
        }
        # checks will be determined after checking all fields against regular expressions and DB

        # field values given by the user for account creation
        mail = request.form.get("mail").strip().replace(" ", "")
        password = request.form.get("password").strip().replace(" ", "")
        conf_pass = (
            request.form.get("conf_pass").strip().replace(" ", "")
        )  # confirm password field
        # the password & confirm password are received in plain-text
        # the leading, trailing & in-between spaces are removed by strip() & replace()

        # will be used as value in case of unsuccessful signup
        values = {"mail": mail}

        # check for mail
        # should not be empty
        # match the regular Expression
        if not mail or not re.fullmatch(mail_regex, mail):
            # if any of the condition fails the
            # valid mail check is marked False
            check["valid_Mail"] = False

        # checking if the mail already exists in the DB
        if collection.count_documents({"mail": "%s" % (mail)}):
            # if user already exists
            check["mail_Duplicate"] = False

        # check for plain-text password & confirm password
        # should not be empty
        # match the regular Expression
        if (
            not password
            or not conf_pass
            or not re.fullmatch(pass_regex, password)
            or not re.fullmatch(pass_regex, conf_pass)
        ):
            # if any of the condition fails the
            # valid password check is marked False
            check["valid_Pass"] = False

        # match password and confirm password fields
        if password != conf_pass:
            # if the passwords do not match
            # match password check is marked False
            check["match_Pass"] = False

        # validation of CAPTCHA
        if not captcha.validate():
            # if incorrect captcha
            check["valid_Captcha"] = False

        if all([i for i in check.values()]):
            ph = ar.PasswordHasher(hash_len=64)

            pepper1, pepper2 = mail.split("@")

            password = ph.hash(pepper1 + password + pepper2)

            print(pepper1)
            print(pepper2)
            print(password)

            # dictionary used as query for creating user in DB
            # and return same values to form in case there is an error
            details = {
                "mail": "%s" % (mail),
                "pass": "%s" % (password),
                "attempts": 3,
                "verified": 0,
            }

            # inersting the details in DB
            # creating the user account
            if collection.insert_one(details):
                # mail id will be stored in token session
                # to keep track of file uploads
                session["tmp_mail"] = mail

                collection = db["code_count"]

                collection.insert_one({"mail": mail, "code": "0", "attempts": 3})

                return redirect("/code_generate")

            else:
                return render_template(
                    "signup.html", check=check, msg=msg, values=values
                )
        else:
            return render_template("signup.html", check=check, msg=msg, values=values)

    return redirect("/signup")


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


@app.route("/file_upload_form", methods=["POST", "GET"])
def file_upload_form():
    file_check = {
        "valid_address": True,
        "valid_area": True,
        "valid_state": True,
        "valid_district": True,
        "valid_zip": True,
        "valid_file": True,
        "valid_Captcha": True,
    }
    msg = ""
    file_values = {"address": "", "area": "", "zipcode": ""}

    remaining_attempts = False

    mail = session.get("mail")
    attempts = get_attempts(mail)

    if attempts > 0:
        remaining_attempts = True

    return render_template(
        "file_upload.html",
        msg="Logged in Successfully",
        check=file_check,
        err="",
        attempts=attempts,
        remaining_attempts=remaining_attempts,
        values=file_values,
    )


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# function to validate file upload forms data
@app.route("/file_upload", methods=["POST", "GET"])
def file_upload():
    check = {
        "valid_address": True,
        "valid_area": True,
        "valid_state": True,
        "valid_district": True,
        "valid_zip": True,
        "valid_file": True,
        "valid_Captcha": True,
    }
    values = ({"address": "", "area": "", "zipcode": ""},)
    # checks will be determined after verifying every field

    # check for session started
    mail = session.get("mail")
    # print("-----------------------------------")
    # print("Session Mail : ", mail, ":::", session.get("mail"))
    # print("-----------------------------------")

    remaining_attempts = False

    attempts = get_attempts(mail)

    if attempts > 0:
        remaining_attempts = True

    address_regex = r"^[0-9a-zA-Z]([0-9a-zA-Z\, \-\.]){1,127}"

    area_regex = r"^[0-9a-zA-Z]([0-9a-zA-Z\, \-\.]){1,63}"

    # after the file upload form submission
    # All POST requests get validated here ↓
    if request.method == "POST":
        # field values given by the user for image location
        address = request.form.get("address").strip()
        area = request.form.get("area").strip()
        state = request.form.get("state").strip().replace(" ", "")
        district = request.form.get("district").strip().replace(" ", "")
        zipcode = request.form.get("zipcode").strip().replace(" ", "")
        # the leading, trailing & in-between spaces are removed by strip() & replace()

        # the file uploaded by user
        file = request.files["road_img"]

        # replaces any extra characters with underscore (_)
        file_name = secure_filename(file.filename)
        print(file_name)

        # result variables initialized with 0 value
        pothole_percent, detect_result = 0, 0

        # validation of CAPTCHA
        # print('Answer : ', captcha.get_answer(), captcha.validate(), captcha)
        if not captcha.validate():
            check["valid_Captcha"] = False

        # if the file name matches the allowed extensions
        if allowed_file(file_name):
            # save the file for pothole detection at uploads folder
            file.save(os.path.join(app.config["UPLOAD_FOLDER"], file_name))

            # file path is stored that will be fed to ML model for pothole detection
            path = os.path.join(app.config["UPLOAD_FOLDER"], file_name)

            if valid_file(UPLOAD_FOLDER + file_name):
                # pothole detection results after ML model
                pothole_percent, detect_result = model(glob.glob(path))

            else:
                # invalid file type
                check["valid_file"] = False

        else:
            # invalid file type
            check["valid_file"] = False

        if not address or not re.fullmatch(address_regex, address):
            # empty address field
            check["valid_address"] = False

        if not area or not re.fullmatch(area_regex, area):
            # empty area field
            check["valid_area"] = False

        # list of states valid for form submission
        STATES_LIST_INDIA = ["Gujarat", "Maharashtra", "Delhi"]

        if not state or state not in STATES_LIST_INDIA:
            # state field either empty or not in list
            check["valid_state"] = False

        if state == "Gujarat":
            districts = [
                "Ahmedabad",
                "Bhavnagar",
                "Gandhinagar",
                "Rajkot",
                "Surat",
                "Vadodara",
            ]
        elif state == "Maharashtra":
            districts = [
                "Mumbai City",
                "Mumbai Suburban",
                "Nagpur",
                "Nashik",
                "Pune",
            ]
        elif state == "Delhi":
            districts = [
                "Central Delhi",
                "East Delhi",
                "New Delhi",
                "North Delhi",
                "South Delhi",
                "West Delhi",
            ]
        else:
            districts = []
            check["valid_state"] = False

        if not district or district not in districts:
            # empty district field
            check["valid_district"] = False

        if (
            not zipcode
            or zipcode.startswith("0")
            or len(zipcode) != 6
            or not zipcode.isnumeric()
        ):
            # empty or invalid zipcode
            check["valid_zip"] = False

        parameters = {
            "address": address,
            "area": area,
            "state": state,
            "district": district,
            "zipcode": zipcode,
            "file_name": file_name,
        }

        results = {"pothole_percent": pothole_percent, "detect_result": detect_result}

        values = {"address": address, "area": area, "zipcode": zipcode}

        remaining_attempts = False

        attempts = get_attempts(mail)
        # print(mail)
        # print(check)

        if attempts > 0:
            remaining_attempts = True

        if all([i for i in check.values()]):
            mail = session.get("mail")

            db = mongoClient["userdata"]

            collection = db["pothole_result"]

            details = {
                "mail": "%s" % mail,
                "address": "%s" % address,
                "area": "%s" % area,
                "state": "%s" % state,
                "district": "%s" % district,
                "zipcode": "%s" % zipcode,
                "file_name": "%s" % file_name,
                "pothole_percent": "%s" % pothole_percent,
                "timestamp": "%s" % dt.now(),
            }

            if pothole_percent < 60:
                os.remove(UPLOAD_FOLDER + file_name)

                # print(session.get("mail"))

                # print(list(db["data"].find({"mail": mail})))

                db["data"].update_many({"mail": mail}, {"$inc": {"attempts": -1}})

                # print(list(db["data"].find({"mail": mail})))

                remaining_attempts = False

                attempts = get_attempts(mail)

                if attempts > 0:
                    remaining_attempts = True

                return render_template(
                    "pothole_result.html",
                    parameters=parameters,
                    results=results,
                    attempts=attempts,
                    remaining_attempts=remaining_attempts,
                )

            else:
                if collection.insert_one(details):
                    # print("inserted", details)

                    # print(session.get("mail"))

                    # print(list(db["data"].find({"mail": mail})))

                    db["data"].update_many({"mail": mail}, {"$inc": {"attempts": -1}})

                    # print(list(db["data"].find({"mail": mail})))

                    remaining_attempts = False

                    attempts = get_attempts(mail)

                    if attempts > 0:
                        remaining_attempts = True

                    return render_template(
                        "pothole_result.html",
                        parameters=parameters,
                        results=results,
                        attempts=attempts,
                        remaining_attempts=remaining_attempts,
                    )
                else:
                    return render_template(
                        "file_upload.html",
                        check=check,
                        msg="",
                        err="Some Error Occurred. Try Again.",
                        attempts=attempts,
                        remaining_attempts=remaining_attempts,
                        values=values,
                    )
        else:
            # if session["mail"]:
            #     return render_template(
            #         "pothole_result.html",
            #         parameters=parameters,
            #         results=results,
            #         attempts=attempts,
            #         remaining_attempts=remaining_attempts,
            #     )

            return render_template(
                "file_upload.html",
                check=check,
                msg="",
                err="",
                attempts=attempts,
                remaining_attempts=remaining_attempts,
                values=values,
            )

    return redirect("/file_upload_form")


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# ADMIN SECTION

# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

# A D M I N
# 			A D M I N
# 					  A D M I N


# Admin Login Form
@app.route("/admin")
def admin():
    errors = {
        "valid_ID": True,
        "valid_Pass": True,
        "valid_Captcha": True,
        "valid_admin": True,
        # error checks will be marked after the validations
    }
    # load admin login form
    return render_template("admin.html", errors=errors)


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# Function to validate admin login and redirect to Admin Dashboard if login success
@app.route("/admin_login", methods=["POST", "GET"])
def admin_login():
    errors = {
        "valid_ID": True,
        "valid_Pass": True,
        "valid_Captcha": True,
        "valid_admin": True,
    }
    # error checks will be marked after the validations
    # after the form is submitted
    if request.method == "POST":
        uid_regex = r"^[a-zA-Z]{3}[0-9]{11}$"
        # Regular Expression for valid admin ID

        pass_regex = (
            r"^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*#?&])[A-Za-z\d@$!#%*?&]{8,20}$"
        )
        # Regular Expression for valid password
        # must include :
        # 1 Uppercase
        # 1 Lowercase
        # 1 Number/Digit
        # 1 Special Character
        # length must be between 8 to 20

        # credentials received in plain-text
        # DB contains hashed_salted form of credentials
        uid = request.form.get("uid").strip().replace(" ", "").upper()
        password = request.form.get("password").strip().replace(" ", "")
        # the password is received in plain-text
        # the leading, trailing & in-between spaces are removed by strip() & replace()

        ph = ar.PasswordHasher(hash_len=64)

        # validation of CAPTCHA
        if not captcha.validate():
            # if incorrect captcha
            errors["valid_Captcha"] = False

        if not uid or not re.fullmatch(uid_regex, uid):
            # if any of the condition fails the
            # valid ID check is marked False
            errors["valid_ID"] = False

        if not password or not re.fullmatch(pass_regex, password):
            # if any of the condition fails the
            # valid password check is marked False
            errors["valid_Pass"] = False

        # database to be used for verifying Login info
        db = mongoClient["userdata"]

        # the collection used for valid user credentials
        collection = db["admin"]

        # since all the credentials are hashed we need to get all the records of DB
        # check each id-password pair
        all_doc = collection.find({}, {"_id": 0, "id": 1, "pass": 1})

        # valid admin check is marked false
        # true if the id-password pair exists
        errors["valid_admin"] = False

        # searching each record
        for document in all_doc:
            # print(document)
            # print(type(document))

            # check if id-password both fields exist
            if len(document) == 2:
                # looping through each record fields to validate credentials

                valid_ID, valid_pass = 0, 0
                for field, value in document.items():
                    # credentials hashed and salted using Bcrypt
                    # checkpw function used to check the credentials
                    # the function takes credentials(input) & credentials stored in database
                    # both parameters need to be in bytes and thus are encoded
                    # boolean value is returned (True for match)

                    if field == "id":
                        # valid_ID = bcrypt.checkpw(uid.encode(), value.encode())
                        try:
                            valid_ID = ph.verify(value, uid)
                        except ar.exceptions.VerifyMismatchError:
                            pass
                        # id will be true if it matches

                        # print(valid_ID)
                    elif field == "pass":
                        # valid_pass = bcrypt.checkpw(password.encode(), value.encode())
                        try:
                            valid_pass = ph.verify(value, password)
                        except ar.exceptions.VerifyMismatchError:
                            pass
                        # password will be true if it matches

                        # print(valid_pass)

                # if both credentials are correct as pair
                if valid_ID and valid_pass:
                    errors["valid_admin"] = True
                    print(document)
                    break

        print(errors)

        # if all the checks are True
        if all([i for i in errors.values()]):
            # get the graph data in JSON form
            # will be used to plot interactive graph
            return redirect("/admin_dashboard")
            bar_state, bar_state_district = get_counts()

            # redirect to main dashboard and display the graphs
            return render_template(
                "admin_dashboard.html",
                state_bar_graph=bar_state,
                state_district_bar_graph=bar_state_district,
            )

        else:
            return render_template("admin.html", errors=errors)
    # bar_state, bar_state_district = get_counts()

    # return render_template(
    #     "admin_dashboard.html",
    #     state_bar_graph=bar_state,
    #     state_district_bar_graph=bar_state_district,
    # )

    else:
        # return to admin login form
        return redirect("/admin")

    return redirect("/admin")
    # return render_template("admin.html", errors=errors)


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


@app.route("/admin_dashboard", methods=["POST", "GET"])
def admin_dashboard():
    bar_state, bar_state_district = get_counts()

    # redirect to main dashboard and display the graphs
    return render_template(
        "admin_dashboard.html",
        state_bar_graph=bar_state,
        state_district_bar_graph=bar_state_district,
    )


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


@app.errorhandler(CSRFError)
def handle_csrf_error(e):
    return render_template("invalid_csrf.html", reason=e.description), 400


# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -


# MAIN FUNCTION

if __name__ == "__main__":
    app.run(debug=True)

    # -------------------------------------------------------

    # Uncomment after testing

    # file_upload_attempt_update.add_job(
    #     id="Update Attempts", func=update_attempts, trigger="cron", hour=0, minute=0
    # )

    # file_upload_attempt_update.start()

    # clr_unverified_accounts.add_job(
    #     id="Clear Unverified Accounts",
    #     func=clear_unverified_accounts,
    #     trigger="cron",
    #     hour=0,
    #     minute=17,
    # )

    # clr_unverified_accounts.start()

    # verification_code_attempt_update.add_job(
    #     id="Update verification Code Attempts",
    #     func=update_verification_code_attempts,
    #     trigger="cron",
    #     hour=0,
    #     minute=17,
    # )

    # verification_code_attempt_update.start()

    # clr_sessions.add_job(
    #     id="Clear Sessions", func=clear_sessions, trigger="interval", minutes=10
    # )

    # clr_sessions.start()

    # authority_report.add_job(
    #     id="Report Authority",
    #     func=report_authority,
    #     trigger="cron",
    #     week="*",
    #     day_of_week="mon",
    #     hour=0,
    #     minute=0,
    # )

    # authority_report.start()

    # app.run(debug=True, host="0.0.0.0", use_reloader=False)

    # -------------------------------------------------------

    # file_upload_attempt_update.add_job(
    #     id="Update Attempts", func=update_attempts, trigger="cron", hour=0, minute=0
    # )

    # file_upload_attempt_update.start()

    # clr_unverified_accounts.add_job(
    #     id="Clear Unverified Accounts",
    #     func=clear_unverified_accounts,
    #     trigger="cron",
    #     hour=0,
    #     minute=17,
    # )

    # clr_unverified_accounts.start()

    # verification_code_attempt_update.add_job(
    #     id="Update verification Code Attempts",
    #     func=update_verification_code_attempts,
    #     trigger="cron",
    #     hour=0,
    #     minute=17,
    # )

    # verification_code_attempt_update.start()

    # clr_sessions.add_job(
    #     id="Clear Sessions", func=clear_sessions, trigger="interval", minutes=10
    # )

    # clr_sessions.start()

    # authority_report.add_job(
    #     id="Report Authority",
    #     func=report_authority,
    #     trigger="cron",
    #     week="*",
    #     day_of_week="mon",
    #     hour=0,
    #     minute=0,
    # )

    # authority_report.start()

    # app.run(debug=True, host="0.0.0.0", use_reloader=False)
